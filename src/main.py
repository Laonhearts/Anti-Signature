import argparse
import time
import sys
import hashlib
import os
import re
import mysql.connector  # MySQL 연결용
import cx_Oracle  # Oracle 연결용
from scapy.all import sniff, wrpcap, IP, TCP, UDP, ICMP
import threading
import signal
from datetime import datetime
import shutil
from docx import Document
from datetime import datetime
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

# 파일 시그니처 정의
FILE_SIGNATURES = {
    
    'pdf': b'%PDF',
    'gif': [b'GIF87a', b'GIF89a'],
    'png': b'\x89PNG\r\n\x1a\n',
    'jpg': [b'\xff\xd8\xff\xe0', b'\xff\xd8\xff\xe1', b'\xff\xd8\xff\xe8', b'\xff\xd8\xff\xdb', b'\xff\xd8\xff\xee'],
    'zip': b'PK\x03\x04',
    'exe': b'MZ',
    'msi': b'MZ',  
    'ico': b'\x00\x00\x01\x00',
    'cur': b'\x00\x00\x02\x00',
    'mpg': b'\x00\x00\x01\xb3',
    'doc': [b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1', b'\xec\xa5\xc1\x00', b'\xbe\x00\x00\x00\xab\x00\x00\x00'],
    'xls': [b'\xfd\xff\xff\xff', b'\xfe\xff'],
    'ppt': [b'\x00\x6e\x1e\xf0', b'\xfd\xff\xff\xff'],
    'mp4': b'\x00\x00\x00\x18ftyp',
    'mov': b'moov',
    'bmp': b'BM',
    'tar': b'ustar',
    'gz': [b'\x1f\x8b\x08', b'\x1f\x9d', b'\x1f\xa0'],
    'avi': b'RIFF',
    'wav': b'RIFF',
    'mp3': [b'ID3', b'\xff\xfb'],
    'psd': b'8BPS',
    'rtf': b'{\\rtf',
    'xml': b'<?xml',
    'json': b'{',
    'flv': b'FLV',
    'rm': b'.RMF',
    'tif': [b'II*\x00', b'MM\x00*'],
    'arj': b'\x60\xea',
    'rar': b'Rar!',
    '3gp': [b'\x00\x00\x00\x14ftyp3gp', b'\x00\x00\x00\x20ftyp3g2'],
    'aac': b'ADIF',
    'amr': b'#!AMR',
    'iso': b'CD001',
    'lha': [b'\x2D\x6C\x68', b'\x2D\x6C\x68\x35'],
    'eps': b'%!PS-Adobe',
    'fli': [b'\xAF\x11', b'\x01\x11\xAF'],
    'qxd': [b'\x00\x00\x49\x49\x58\x50\x52', b'\x00\x00\x4D\x4D\x58\x50\x52'],
    'ai': b'%!PS-Adobe',
    'wma': b'0&\xb2u',
    'wmv': b'0&\xb2u',
    'asf': b'0&\xb2u',
    '7z': b'7z\xbc\xaf\x27\x1c',
    'bz2': b'BZh',
    'cab': b'MSCF',
    'dmg': b'x',
    'jar': b'PK\x03\x04',
    'gzip': b'\x1F\x8B',
    'arc': b'\x1A',
    'jpeg': [b'\xFF\xD8\xFF\xDB', b'\xFF\xD8\xFF\xEE', b'\xFF\xD8\xFF\xE1'],
    'pif': b'\x00',
    'mac': b'\x00\x00\x00\x02',
    'wmf': b'\xd7\xcd\xc6\x9a',
    'mp2': b'\xff\xf3',
    'dbf': [b'\x02', b'\x03'],
    'db': b'\x00\x01',
    'mdb': b'\x00\x01\x00\x00Standard Jet DB',
    'emf': b'\x01\x00\x00\x00',
    'docx': b'PK\x03\x04',
    'pptx': b'PK\x03\x04',
    'xlsx': b'PK\x03\x04',

}

# 랜섬웨어와 관련된 확장자 목록
RANSOMWARE_EXTENSIONS = [
    # 랜섬웨어로 의심되는 확장자 목록

    'locky', 'zepto', 'odin', 'cerber', 'crysis', 'wallet', 'zzzzz', 'ccc', 'exx', 'ecc', 'crypt', 'crab', 'cbf',
    'arena', 'dharma', 'arrow', 'grt', 'ryuk', 'phobos', 'krab', 'fucked', 'crypted', 'satan', 'xrat', 'gandcrab', 
    'cyborg', 'salus', 'ciphered', 'shiva', 'stupid', 'why', 'weapologize', 'grandcrab', 'megacortex', 'revil', 
    'ekvf', 'fairytail', 'exorcist', 'mamba', 'killing_time', 'recovery', 'volcano', 'pscrypt', 'tcps', 'fuxsocy', 
    'heof', 'matrix', 'medusa', 'snake', 'ftcode', 'calipso', 'encrypted', 'vault', 'shariz', 'tor', 'globeimposter', 
    'kogda', 'crypmic', 'bokbot', 'dewar', 'defray', 'greystar', 'pclock', 'moisha', 'zcrypt', 'djvu', 'dotmap', 
    'sodinokibi', 'calipto', 'fun', 'worm', 'inferno', 'medy', 'fudcrypt', 'zohar', 'nig', 'ogre', 'uizsdgpo', 
    'hermes', 'nellyson', 'pay', 'bepow', 'futc', 'suck', 'f**ked', 'foxf***r', 'hotjob', 'gusar', 'jigsaw', 'oklah', 
    'lockfile', 'aesc', 'wb', 'tycoon', 'careto', 'azov', 'conti', 'ragnarok', 'crippled', 'trosy', 'johnny', 'regen', 
    'whatthe', '000', 'hitler', 'somethinghere', 'uw', 'anubi', 'notgotenough', 'justforyou', 'bondy', 'ark', 'kubos', 
    'police', 'aprilfool', 'bitch', 'crisis', 'teerac', 'long', 'infom', 'woswos', 'uw4w', 'ctrb', 'purge', 'fiwlgphz', 
    'platano', 'pro', 'locked', 'eyy', 'whythis', 'k0der', 'joker', 'fucku', 'charm', 'doubleup', 'payup', 'firecrypt', 
    'miyake', 'senpai', 'onspeed', 'fxckoff', 'zer0', 'p0rn', 'beethoven', 'im sorry', 'kuba', 'hahaha', '555', 
    'sexx', 'fuckyou', 'supercrypt', 'gudrotka', 'catchmeifyoucan', 'bambam', 'lucy', 'sadism', 'fedup', 'makop', 
    'alpha', 'master', 'mcafee', 'bastard', 'locki', 'striker', 'grimly', 'cryptolocker', 'kukaracha', 'kraken', 
    'supersystem', 'hot', 'ispy', 'newversion', 'payfast', 'futurat', 'unlock', 'kkk', 'openme', 'blablabla', 
    'goof', 'psycho', 'trigger', 'memeware', 'emotet', 'wannacry', 'notpetya', 'mazefuck', 'mazelocker', 'pegasus', 
    'sodinokibi', 'rdp', 'kodg', 'covm', 'cazw', 'egregor', 'revil', 'recovery', 'sodin', 'kodc', 'gdc', 'gdcb', 
    'grb', 'egregor', 'medusalocker', 'medusa', 'phobos', 'ransom', 'makop', 'mountlocker', 'avoslocker', 'cuba', 
    'kaseya', 'kaseyacrypt', 'blackcat', 'alphv', 'lorenz', 'lockbit', 'abcd', 'lukitus', 'moqs', 'thunderx'

]

# 악성 패킷 탐지 패턴 목록
MALICIOUS_PATTERNS = [
    
    {"proto": "TCP", "dport": 4444},  # 예: 일반적으로 악성 활동에 사용되는 포트
    {"proto": "TCP", "flags": "FPU"},  # 예: 비정상적인 플래그 조합
    {"proto": "UDP", "dport": 53413},  # 예: UDP에서 흔히 사용되는 악성 포트
    {"proto": "ICMP", "type": 8, "code": 0},  # 예: ICMP Echo Request 공격
    {"proto": "TCP", "dport": 8080},  # 8080 포트에서의 TCP 통신 감지
    {"proto": "UDP", "sport": 12345}  # 12345 포트에서의 UDP 통신 감지

]

# 네트워크 패킷 캡처와 악성 패킷 필터링 기능
stop_sniffing = False  # 패킷 캡처 중지 플래그
captured_packets = []  # 캡처된 패킷을 저장할 리스트
pcap_file = 'packets.pcap'  # 패킷 저장 파일

CANARY_VALUE = b'ANTISIG'  # 카나리 값 (특정한 값으로 설정)

# DB 연결 객체
db_conn = None
cursor = None

def print_ascii_art():  # 프로그램 시작시 표기될 아스키아트 (ANTI SIGNATURE)

    ascii_art = """

    █████╗ ███╗   ██╗████████╗██╗    ███████╗██╗ ██████╗ ███╗   ██╗ █████╗ ████████╗██╗   ██╗██████╗ ███████╗
    ██╔══██╗████╗  ██║╚══██╔══╝██║    ██╔════╝██║██╔════╝ ████╗  ██║██╔══██╗╚══██╔══╝██║   ██║██╔══██╗██╔════╝
    ███████║██╔██╗ ██║   ██║   ██║    ███████╗██║██║  ███╗██╔██╗ ██║███████║   ██║   ██║   ██║██████╔╝█████╗  
    ██╔══██║██║╚██╗██║   ██║   ██║    ╚════██║██║██║   ██║██║╚██╗██║██╔══██║   ██║   ██║   ██║██╔══██╗██╔══╝  
    ██║  ██║██║ ╚████║   ██║   ██║    ███████║██║╚██████╔╝██║ ╚████║██║  ██║   ██║   ╚██████╔╝██║  ██║███████╗
    ╚═╝  ╚═╝╚═╝  ╚═══╝   ╚═╝   ╚═╝    ╚══════╝╚═╝ ╚═════╝ ╚═╝  ╚═══╝╚═╝  ╚═╝   ╚═╝    ╚═════╝ ╚═╝  ╚═╝╚══════╝
    
    """
    
    print(ascii_art)

def show_loading_effect():  # 로딩 효과 출력

    for _ in range(10):

        sys.stdout.write(". ")    
        sys.stdout.flush()
        
        time.sleep(0.5)
    
    print()

def calculate_file_hash(file_path, hash_algorithm='sha256'):    # 파일 해시 계산
    
    hash_func = hashlib.new(hash_algorithm)
    
    try:
    
        with open(file_path, 'rb') as f:
    
            while chunk := f.read(8192):
    
                hash_func.update(chunk)
    
    except FileNotFoundError:
    
        print(f"Error: {file_path} 파일을 찾을 수 없습니다.")
    
        return None

    return hash_func.hexdigest()

def connect_to_db(db_type):

    global db_conn, cursor
    
    if db_type == 'mysql':
    
        # MySQL 연결 설정
        db_conn = mysql.connector.connect(
    
            host="localhost",
            user="root",
            password="",
            database="signature",
            port=3306
    
        )
    
        cursor = db_conn.cursor()
    
        print(f"{db_type.upper()} 데이터베이스에 연결되었습니다.")

    elif db_type == 'oracle':
    
        # Oracle 연결 설정
    
        dsn = cx_Oracle.makedsn("localhost", 1521, service_name="orcl")
    
        db_conn = cx_Oracle.connect(user="system", password="2558jun@", dsn=dsn)
    
        cursor = db_conn.cursor()
    
        print(f"{db_type.upper()} 데이터베이스에 연결되었습니다.")
    
def reconnect_db():
    
    if not db_conn or not db_conn.is_connected():
    
        print("DB 연결이 끊어졌습니다. 재연결을 시도합니다.")
    
        connect_to_db('mysql')
    
def insert_operation_log(operation, details, status):
    
    if cursor:
    
        timestamp = datetime.now()

        # MySQL 쿼리: 자리 표시자를 %s로 변경
        query = "INSERT INTO operation_logs (operation, details, status, timestamp) VALUES (%s, %s, %s, %s)"
        values = (operation, details, status, timestamp)

        # 쿼리와 값을 출력하여 디버깅
        print(f"Executing query: {query}")
        print(f"With values: {values}")

        cursor.execute(query, values)
        db_conn.commit()

        print(f"로그가 저장되었습니다: {operation}, {details}, {status}, {timestamp}")


def insert_file_integrity_log(file_name, action, status):   # 파일 무결성, 랜섬웨어 감염 여부, 안티디버깅 등과 관련된 로그를 DB에 삽입하는 함수
    
    if cursor:
        
        timestamp = datetime.now()
        
        if isinstance(db_conn, mysql.connector.connection.MySQLConnection):
        
            query = "INSERT INTO file_integrity_logs (file_name, action, status, timestamp) VALUES (%s, %s, %s, %s)"
        
            values = (file_name, action, status, timestamp)
        
        else:
        
            query = "INSERT INTO file_integrity_logs (file_name, action, status, timestamp) VALUES (:1, :2, :3, :4)"
        
            values = (file_name, action, status, timestamp)
        
        cursor.execute(query, values)
        
        db_conn.commit()
        
        print(f"파일 무결성 로그가 저장되었습니다: {file_name}, {action}, {status}, {timestamp}")

def insert_file_signature_log(file_name, signature_before, signature_after): # 파일 시그니처 변경 로그를 DB에 삽입하는 함수
    
    if cursor:
        
        timestamp = datetime.now()
        
        if isinstance(db_conn, mysql.connector.connection.MySQLConnection):
        
            query = "INSERT INTO file_signature_logs (file_name, signature_before, signature_after, timestamp) VALUES (%s, %s, %s, %s)"
        
            values = (file_name, signature_before, signature_after, timestamp)
        
        else:
        
            query = "INSERT INTO file_signature_logs (file_name, signature_before, signature_after, timestamp) VALUES (:1, :2, :3, :4)"
        
            values = (file_name, signature_before, signature_after, timestamp)
        
        cursor.execute(query, values)
        
        db_conn.commit()
        
        print(f"파일 시그니처 로그가 저장되었습니다: {file_name}, {signature_before}, {signature_after}, {timestamp}")

def check_file_integrity(file_path, expected_hash=None):
    
    try:

        with open(file_path, 'rb') as f:

            file_signature = f.read(20)

    except FileNotFoundError:

        print(f"Error: {file_path} 파일을 찾을 수 없습니다.")

        insert_operation_log("File Integrity Check", file_path, "File Not Found")

        return

    file_extension = file_path.split('.')[-1].lower()

    suspicious = False

    if file_extension in FILE_SIGNATURES:

        expected_signatures = FILE_SIGNATURES[file_extension]

        if not isinstance(expected_signatures, list):

            expected_signatures = [expected_signatures]
        
        if any(file_signature.startswith(sig) for sig in expected_signatures):

            print(f"{file_path}의 파일 시그니처가 정상입니다.")

            insert_operation_log("File Integrity Check", file_path, "Passed")

        else:

            print(f"경고: {file_path}의 파일 시그니처가 예상과 다릅니다.")

            suspicious = True

            insert_operation_log("File Integrity Check", file_path, "Signature Mismatch")

    else:

        print(f"알 수 없는 확장자입니다: {file_extension}")

        suspicious = True

        insert_operation_log("File Integrity Check", file_path, "Unknown Extension")

    if suspicious:

        insert_operation_log("File Integrity Check", file_path, "Suspicious")

    if expected_hash:

        calculated_hash = calculate_file_hash(file_path)

        if calculated_hash and calculated_hash != expected_hash:

            print(f"경고: {file_path}의 파일 해시 무결성이 훼손되었습니다.")

            insert_operation_log("File Integrity Check", file_path, "Hash Mismatch")

        elif calculated_hash:

            print(f"{file_path}의 파일 해시가 무결합니다.")

            insert_operation_log("File Integrity Check", file_path, "Hash Match")



def check_for_ransomware(file_path):

    file_name = file_path.split('/')[-1].lower()

    suspicious = False

    if any(ext in file_name for ext in RANSOMWARE_EXTENSIONS):

        print(f"경고: {file_path}는 랜섬웨어와 관련된 확장자를 포함하고 있습니다.")

        suspicious = True

        insert_operation_log("Ransomware Detection", file_path, "Suspicious Extension")

    if "readme" in file_name or "decrypt" in file_name:

        print(f"경고: {file_path}는 랜섬웨어 관련 파일일 수 있습니다.")

        suspicious = True

        insert_operation_log("Ransomware Detection", file_path, "Suspicious File Name")

    if not suspicious:

        print(f"{file_path}는 랜섬웨어에 감염되지 않은 정상 파일입니다.")

        insert_operation_log("Ransomware Detection", file_path, "Clean")

    return suspicious

def apply_anti_debugging_and_obfuscation(file_path):  # PE 또는 ELF 파일에 안티디버깅 및 난독화 기법을 적용하는 함수
    
    if not os.path.exists(file_path):
    
        print(f"Error: {file_path} 파일을 찾을 수 없습니다.")
    
        return
    print(f"{file_path}에 안티디버깅 및 난독화 기법을 적용 중입니다...")
    
    anti_debugging_code = b"\xEB\xFE"  # 무한 루프 삽입
    
    with open(file_path, 'ab') as f:
    
        f.write(anti_debugging_code)
    
    with open(file_path, 'ab') as f:
    
        random_bytes = os.urandom(10)
    
        f.write(random_bytes)
    
    print("안티디버깅 및 난독화 기법이 성공적으로 적용되었습니다.")

def detect_anti_debugging_and_obfuscation(file_path):  # 파일에 안티디버깅 및 난독화 기법이 적용되었는지 확인하는 함수
    
    if not os.path.exists(file_path):
    
        print(f"Error: {file_path} 파일을 찾을 수 없습니다.")
    
        return
    
    print(f"{file_path}에 안티디버깅 및 난독화가 적용되었는지 확인 중입니다...")
    
    with open(file_path, 'rb') as f:
    
        content = f.read()
    
        if b"\xEB\xFE" in content:
    
            print(f"{file_path}에 안티디버깅 기법이 적용되었습니다.")
    
        else:
    
            print(f"{file_path}에 안티디버깅 기법이 적용되지 않았습니다.")
            
def remove_anti_debugging_and_obfuscation(file_path):   # 안티디버깅 기법과 난독화 기법을 해제하는 함수
    
    if not os.path.exists(file_path):
    
        print(f"Error: {file_path} 파일을 찾을 수 없습니다.")
    
        return

    print(f"{file_path}에서 안티디버깅 및 난독화 기법을 제거하는 중입니다...")

    # 파일에서 안티디버깅 코드 및 난독화된 바이트 패턴을 찾고 제거
    with open(file_path, 'rb') as f:
    
        content = f.read()

    # 안티디버깅 코드 제거
    content = content.replace(b"\xEB\xFE", b"")

    # 파일을 덮어쓰기 모드로 열고 변경된 내용을 기록
    with open(file_path, 'wb') as f:
    
        f.write(content)

    print("안티디버깅 및 난독화 기법이 성공적으로 제거되었습니다.")
    
def insert_canary(file_path):  # 파일에 카나리 값을 삽입하는 함수
    
    print(f"파일에 삽입된 카나리 값: {CANARY_VALUE.hex()}")
    
    try:
    
        with open(file_path, 'rb+') as f:
    
            f.seek(0, os.SEEK_END)  # 파일 끝으로 이동
            f.write(b'\x00' * 16)   # 빈 공간 확보 (예제: 16바이트)
            f.write(CANARY_VALUE)   # 카나리 값 삽입
    
            print(f"카나리 값이 {file_path}에 성공적으로 삽입되었습니다.")
    
    except FileNotFoundError:
    
        print(f"Error: {file_path} 파일을 찾을 수 없습니다.")

def check_canary_integrity(file_path):  # 파일의 카나리 값 무결성을 체크하는 함수
    
    try:
    
        with open(file_path, 'rb') as f:
    
            f.seek(-len(CANARY_VALUE), os.SEEK_END)  # 파일 끝에서 카나리 값 읽기
    
            stored_canary = f.read(len(CANARY_VALUE))
    
            print(f"파일에서 읽은 카나리 값: {stored_canary.decode()}")

            # 파일의 실제 카나리 값을 검사
            if stored_canary == CANARY_VALUE:
    
                print("카나리 무결성 검증 통과: 파일이 변조되지 않았습니다.")
    
            else:
    
                print("경고: 카나리 무결성 검증 실패! 파일이 변조되었을 수 있습니다.")
    
    except FileNotFoundError:
    
        print(f"Error: {file_path} 파일을 찾을 수 없습니다.")
    
    except Exception as e:
    
        print(f"Error during canary integrity check: {e}")
        
def remove_canary(file_path):  # 파일에서 카나리 값을 제거하는 함수
    
    try:
    
        with open(file_path, 'rb+') as f:
    
            f.seek(-len(CANARY_VALUE), os.SEEK_END)  # 파일 끝에서 카나리 값 읽기 위치로 이동
    
            stored_canary = f.read(len(CANARY_VALUE))

            if stored_canary == CANARY_VALUE:  # 카나리 값이 일치하는지 확인
    
                print(f"카나리 값이 {file_path}에서 감지되었습니다. 제거 중입니다...")
    
                f.seek(-len(CANARY_VALUE), os.SEEK_END)  # 카나리 위치로 다시 이동
                f.truncate(f.tell())  # 파일을 카나리 값 바로 앞에서 자르기
    
                print("카나리 값이 성공적으로 제거되었습니다.")
    
            else:
    
                print("카나리가 파일에 존재하지 않거나 손상되었습니다.")
    
    except FileNotFoundError:
    
        print(f"Error: {file_path} 파일을 찾을 수 없습니다.")
    
    except Exception as e:
    
        print(f"Error during canary removal: {e}")
        

def process_file_with_cp_option(file_path):  # cp 옵션 실행 함수
    
    temp_folder = 'temp'

    # temp 폴더 생성
    if not os.path.exists(temp_folder):
    
        os.makedirs(temp_folder)

    # 파일 확장자 추출
    original_extension = os.path.splitext(file_path)[1]

    # 확장자 정보를 temp 폴더에 txt 파일로 저장
    extension_info_file = os.path.join(temp_folder, os.path.basename(file_path) + "_extension.txt")
    
    with open(extension_info_file, 'w') as f:
    
        f.write(original_extension)
    
    print(f"원본 파일의 확장자가 {extension_info_file}에 저장되었습니다: {original_extension}")

    # 파일 복사
    copied_file_path = os.path.join(temp_folder, os.path.basename(file_path))
    
    shutil.copy2(file_path, copied_file_path)
    
    print(f"{file_path}가 {copied_file_path}로 복사되었습니다.")

    # 파일 확장자를 .dll로 변경
    new_copied_file_path = os.path.splitext(copied_file_path)[0] + '.dll'
    
    os.rename(copied_file_path, new_copied_file_path)
    
    print(f"복사 파일의 확장자가 {new_copied_file_path}로 변경되었습니다.")

    # 원본 파일에 카나리 삽입
    insert_canary(file_path)

def process_back_option(file_path):  # back 옵션 실행 함수
    
    temp_folder = 'temp'

    # temp 폴더에서 확장자 정보를 읽어오기
    extension_info_file = os.path.join(temp_folder, os.path.basename(file_path) + "_extension.txt")
    
    if not os.path.exists(extension_info_file):
    
        print(f"확장자 정보 파일을 찾을 수 없습니다: {extension_info_file}")
    
        return

    with open(extension_info_file, 'r') as f:
    
        original_extension = f.read().strip()
    
    print(f"원본 파일의 확장자 정보가 {extension_info_file}에서 읽혔습니다: {original_extension}")

    # temp 폴더에서 .dll 파일을 .txt로 되돌림
    copied_file_path = os.path.join(temp_folder, os.path.basename(file_path).replace('.dll', '.txt'))
    
    new_file_path = os.path.join(temp_folder, os.path.splitext(os.path.basename(file_path))[0] + original_extension)
    
    if not os.path.exists(copied_file_path):
    
        print(f"복사된 .dll 파일을 찾을 수 없습니다: {copied_file_path}")
    
        return

    # .dll 확장자에서 .txt로 되돌림
    os.rename(copied_file_path, new_file_path)
    
    print(f"{copied_file_path}가 {new_file_path}로 확장자가 변경되었습니다.")

    # 원본 파일 복원
    exe_file_path = os.path.splitext(file_path)[0] + '.exe'
    
    if os.path.exists(exe_file_path):
    
        final_file_path = os.path.splitext(exe_file_path)[0] + original_extension
    
        os.rename(exe_file_path, final_file_path)
    
        print(f"{exe_file_path}가 {final_file_path}로 원본 확장자로 복원되었습니다.")
    
    else:
    
        print(f"변경된 exe 파일을 찾을 수 없습니다: {exe_file_path}")

def insert_canary(file_path):  # 파일에 카나리 값을 삽입하는 함수
    
    CANARY_VALUE = b'ANTISIG'
    
    print(f"파일에 삽입된 카나리 값: {CANARY_VALUE.hex()}")

    try:
    
        with open(file_path, 'rb+') as f:
    
            f.seek(0, os.SEEK_END)  # 파일 끝으로 이동
            f.write(b'\x00' * 16)   # 빈 공간 확보 (예제: 16바이트)
            f.write(CANARY_VALUE)   # 카나리 값 삽입
    
        print(f"카나리 값이 {file_path}에 성공적으로 삽입되었습니다.")
    
    except FileNotFoundError:
    
        print(f"Error: {file_path} 파일을 찾을 수 없습니다.")
    
    print(f"파일에 삽입된 카나리 값: {CANARY_VALUE.hex()}")
    
def signal_handler(sig, frame):
    
    global stop_sniffing
    
    stop_sniffing = True
    
    print("\n패킷 캡처 종료 중...")

def check_packet_for_malicious_activity(packet):

    if packet.haslayer(IP):

        if packet.haslayer(TCP):

            tcp_layer = packet[TCP]

            for pattern in MALICIOUS_PATTERNS:

                if pattern["proto"] == "TCP" and (pattern.get("dport") is None or tcp_layer.dport == pattern["dport"]):

                    if "flags" in pattern and pattern["flags"] in tcp_layer.flags:

                        insert_operation_log("Network Monitoring", "TCP Packet Detected", "Malicious")

                        return True

        if packet.haslayer(UDP):

            udp_layer = packet[UDP]

            for pattern in MALICIOUS_PATTERNS:

                if pattern["proto"] == "UDP" and (pattern.get("dport") is None or udp_layer.dport == pattern["dport"]):

                    insert_operation_log("Network Monitoring", "UDP Packet Detected", "Malicious")

                    return True

        if packet.haslayer(ICMP):

            icmp_layer = packet[ICMP]

            for pattern in MALICIOUS_PATTERNS:

                if pattern["proto"] == "ICMP" and (pattern.get("type") is None or icmp_layer.type == pattern["type"]):

                    insert_operation_log("Network Monitoring", "ICMP Packet Detected", "Malicious")

                    return True

    insert_operation_log("Network Monitoring", "Packet Detected", "Normal")

    return False


def packet_callback(packet):    # 캡처된 각 패킷을 처리하는 콜백 함수.
    
    global captured_packets
    
    captured_packets.append(packet)  # 캡처된 패킷 저장

    if check_packet_for_malicious_activity(packet):
        
        print(f"[ALERT] 악성 의심 패킷 감지: {packet.summary()}")
    
    else:
        
        print(f"[INFO] 정상 패킷: {packet.summary()}")
        
def start_packet_capture():     # 네트워크 패킷을 캡처하고 처리하는 함수.
    
    print("네트워크 패킷 캡처 시작 (Ctrl + C로 종료)...")
    
    try:
    
        sniff(prn=packet_callback, store=0)  # 실시간으로 패킷 캡처
    
    except KeyboardInterrupt:
    
        print("패킷 캡처가 종료되었습니다.")

def monitor_network():  # 네트워크 모니터링 프로세스 함수.
    
    global stop_sniffing

    stop_sniffing = False

    def signal_handler(sig, frame):
    
        global stop_sniffing
    
        stop_sniffing = True
    
        print("\n패킷 캡처 종료 중...")

    # Ctrl + C 입력시 signal handler 등록
    signal.signal(signal.SIGINT, signal_handler)

    # sniff 종료 조건 설정: stop_sniffing이 True가 되면 종료
    sniff(prn=packet_callback, store=0, stop_filter=lambda x: stop_sniffing)

    print(f"캡처된 패킷을 {pcap_file} 파일로 저장 중...")

    if captured_packets:
    
        wrpcap(pcap_file, captured_packets)  # 캡처된 패킷을 pcap 파일로 저장
    
        print(f"{len(captured_packets)}개의 패킷이 저장되었습니다.")
    
    else:
    
        print("저장할 패킷이 없습니다.")
        
def get_extension_from_signature(file_signature):  #  파일 시그니처 정보를 기반으로 적합한 확장자를 반환하는 함수
    
    for ext, sigs in FILE_SIGNATURES.items():
    
        if not isinstance(sigs, list):
    
            sigs = [sigs]  # 시그니처가 리스트가 아닌 경우 리스트로 변환
    
        if any(file_signature.startswith(sig) for sig in sigs):
    
            return ext
    
    return None  # 시그니처와 일치하는 확장자가 없을 경우

def process_replace_option(file_path):  # replace 옵션 실행 함수    # 파일의 시그니처를 분석하여 적절한 확장자로 파일을 변경하는 함수
    
    try:
        
        # 파일 시그니처 읽기
        with open(file_path, 'rb') as f:
        
            file_signature = f.read(20)  # 첫 20바이트를 시그니처로 사용

        # 시그니처 기반 확장자 결정
        detected_extension = get_extension_from_signature(file_signature)

        if detected_extension:
        
            # 파일 확장자 변경
            new_file_path = os.path.splitext(file_path)[0] + '.' + detected_extension
        
            os.rename(file_path, new_file_path)
        
            print(f"파일의 확장자가 {new_file_path}로 변경되었습니다. (시그니처: {detected_extension.upper()})")
        
        else:
        
            print(f"알 수 없는 시그니처입니다. 확장자를 결정할 수 없습니다: {file_signature.hex()}")

    except FileNotFoundError:
        
        print(f"Error: {file_path} 파일을 찾을 수 없습니다.")
        
def fetch_data_from_db():

    reconnect_db()
    
    try:
    
        if isinstance(db_conn, mysql.connector.connection.MySQLConnection):
    
            cursor.execute("SELECT * FROM operation_logs")
            operation_logs = cursor.fetchall()

            cursor.execute("SELECT * FROM file_integrity_logs")
            integrity_logs = cursor.fetchall()

            cursor.execute("SELECT * FROM file_signature_logs")
            signature_logs = cursor.fetchall()

        elif isinstance(db_conn, cx_Oracle.Connection):
    
            cursor.execute("SELECT * FROM operation_logs")
            operation_logs = cursor.fetchall()

            cursor.execute("SELECT * FROM file_integrity_logs")
            integrity_logs = cursor.fetchall()

            cursor.execute("SELECT * FROM file_signature_logs")
            signature_logs = cursor.fetchall()

        return {
    
            "operation_logs": operation_logs,
            "integrity_logs": integrity_logs,
            "signature_logs": signature_logs
    
        }

    except (mysql.connector.Error, cx_Oracle.DatabaseError) as err:

        print(f"DB에서 데이터를 가져오는 중 오류 발생: {err}")

        return None


def generate_docx_report(report_data, output_file):  # .docx 형식으로 보고서를 생성하는 함수
    
    document = Document()

    document.add_heading('Anti Signature Report', 0)
    
    document.add_paragraph(f"Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    document.add_paragraph('')

    # Operation Logs
    document.add_heading('Operation Logs', level=1)

    previous_date = None

    for log in report_data["operation_logs"]:

        log_date = log[4].date()  # 로그의 날짜 부분만 추출
        
        # 날짜가 변경되면 구분줄 추가
        if previous_date != log_date:

            document.add_paragraph(f"Logs for {log_date}:", style='Intense Quote')

            previous_date = log_date

        paragraph = document.add_paragraph()

        run = paragraph.add_run(f"Operation: {log[1]}, Details: {log[2]}, Status: {log[3]}, Timestamp: {log[4]}")

        # 의심되는 로그를 강조 (노란색 배경)
        if log[3].lower() in ["suspicious", "malicious", "signature mismatch", "hash mismatch"]:

            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # 저장
    document.save(output_file)

    print(f".docx 보고서가 {output_file}에 저장되었습니다.")


def generate_html_report(report_data, output_file):  # HTML 형식으로 보고서를 생성하는 함수

    html_content = f"""
    <html>

    <head>

        <title>Anti Signature Report</title>

        <style>

            body {{ font-family: Arial, sans-serif; }}

            .suspicious {{ background-color: yellow; }}

            .malicious {{ background-color: red; color: white; }}

            .log-date {{ border-top: 2px solid black; margin-top: 10px; padding-top: 10px; }}

        </style>

    </head>

    <body>

        <h1>Anti Signature Report</h1>

        <p>Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>

        <h2>Operation Logs</h2>

    """

    previous_date = None

    for log in report_data["operation_logs"]:

        log_date = log[4].date()  # 로그의 날짜 부분만 추출

        # 날짜가 변경되면 구분줄 추가
        if previous_date != log_date:

            html_content += f"<div class='log-date'><strong>Logs for {log_date}:</strong></div>"

            previous_date = log_date

        log_class = ""

        if log[3].lower() in ["suspicious", "malicious", "signature mismatch", "hash mismatch"]:

            log_class = "suspicious" if log[3].lower() == "suspicious" else "malicious"

        html_content += f"""

            <div class="{log_class}">

                Operation: {log[1]}, Details: {log[2]}, Status: {log[3]}, Timestamp: {log[4]}

            </div>

        """

    html_content += "</body></html>"

    with open(output_file, 'w') as f:

        f.write(html_content)

    print(f"HTML 보고서가 {output_file}에 저장되었습니다.")
    
    
def process_report_option(report_format):  #  보고서 생성 옵션 처리

    # DB에서 데이터 가져오기
    report_data = fetch_data_from_db()

    # 보고서 형식에 따른 처리
    if report_format == 'docx':

        output_file = 'anti_signature_report.docx'

        generate_docx_report(report_data, output_file)
        
    elif report_format == 'html':

        output_file = 'anti_signature_report.html'

        generate_html_report(report_data, output_file)

    else:

        print("지원하지 않는 형식입니다. docx 또는 html을 선택하세요.")


def main():
    
    print_ascii_art()
    
    show_loading_effect()

    parser = argparse.ArgumentParser(description="Anti Signature 프로그램: 파일 무결성 검사 도구 및 랜섬웨어 감염 여부 확인 도구")
    
    parser.add_argument(
    
        '-f', '--file', 
        type=str, 
        help='무결성을 검사할 파일 경로를 지정하고, 선택적으로 파일 해시 값을 제공할 수 있습니다. 예: -f example.pdf:d41d8cd98f00b204e9800998ecf8427e'
    
    )
    
    parser.add_argument(
    
        '-replace', '--replace', 
        action='store_true', 
        help='파일의 시그니처를 기반으로 확장자를 변경합니다.'
    
    )
    
    parser.add_argument(
    
        '-R', '--ransomware', 
        action='store_true', 
        help='파일의 랜섬웨어 감염 여부를 확인합니다.'
    
    )
    
    parser.add_argument(
    
        '-D', '--apply-debug', 
        action='store_true', 
        help='PE 또는 ELF 파일에 안티디버깅 및 난독화 기법을 적용합니다.'
    
    )
    
    parser.add_argument(
    
        '-dd', '--detect-debug', 
        action='store_true', 
        help='PE 또는 ELF 파일에 안티디버깅 및 난독화 기법이 적용되었는지 확인합니다.'
    
    )
    
    parser.add_argument(
    
        '-dx', '--remove-debug', 
        action='store_true', 
        help='PE 또는 ELF 파일에서 안티디버깅 및 난독화 기법을 제거합니다.'
    
    )
    
    parser.add_argument(
    
        '-net', '--network-monitor', 
        action='store_true', 
        help='네트워크 모니터링 및 악성 패킷 필터링을 수행합니다.'
    
    )
    
    parser.add_argument(
    
        '-c', '--canary', 
        action='store_true', 
        help='파일 내의 빈 공간에 카나리를 추가합니다.'
    
    )
    
    parser.add_argument(
    
        '-ccheck', '--canary-check', 
        action='store_true', 
        help='파일의 카나리 무결성을 체크합니다.'
    
    )
    
    parser.add_argument(
    
        '-cd', '--canary-delete',
        action='store_true',
        help='파일에서 카나리를 제거합니다.'
    
    )
    
    parser.add_argument(
    
        '-cp', '--copy-and-process',
        action='store_true',
        help='파일을 temp 폴더로 복사한 후 시그니처를 .exe로 변경합니다. (파일 백업)'
    
    )
    
    parser.add_argument(
        
        '-db', '--database', 
        choices=['mysql', 'oracle'], 
        help='DB 선택 (MySQL 또는 Oracle)'
        
    )
    
    parser.add_argument(
    
        '-dbi', '--db-info', 
        action='store_true', 
        help='DB에 로그 정보 저장'
    
    )
    
    parser.add_argument(
        
        '-report', '--report',
        type=str,
        choices=['docx', 'html'],
        help='DB 데이터를 기반으로 보고서를 생성합니다. 지원되는 형식: docx, html'
    
    )
    
    args = parser.parse_args()

    # DB 연결 설정
    if args.database:
    
        connect_to_db(args.database)
    
    # 파일 관련 옵션 처리
    if args.file:
    
        file_info = args.file.split(':')
    
        file_path = file_info[0]
    
        expected_hash = file_info[1] if len(file_info) > 1 else None

        # 파일 무결성 검사
        check_file_integrity(file_path, expected_hash)
        
        if args.replace:
            
            process_replace_option(file_path)

        if args.ransomware:
    
            check_for_ransomware(file_path)
        
        if args.apply_debug:
    
            apply_anti_debugging_and_obfuscation(file_path)
        
        if args.detect_debug:
    
            detect_anti_debugging_and_obfuscation(file_path)
        
        if args.remove_debug:
    
            remove_anti_debugging_and_obfuscation(file_path)
        
        if args.canary:
    
            insert_canary(file_path)
        
        if args.canary_check:
    
            check_canary_integrity(file_path)
        
        if args.canary_delete:
    
            remove_canary(file_path)
        
        if args.copy_and_process:
    
            process_file_with_cp_option(file_path)

        # DB 로그 저장 처리
        if args.db_info and cursor:
    
            # 파일 무결성 관련 로그 저장
            if expected_hash:
    
                actual_hash = calculate_file_hash(file_path)
    
                status = 'Integrity Passed' if actual_hash == expected_hash else 'Integrity Failed'
    
                insert_file_integrity_log(file_path, "File Integrity Check", status)
            
            # 랜섬웨어 관련 로그 저장
            if args.ransomware:
    
                insert_file_integrity_log(file_path, "Ransomware Check", "Suspicious" if check_for_ransomware(file_path) else "Clean")
            
            # 안티디버깅 및 난독화 관련 로그 저장
            if args.apply_debug:
    
                insert_file_integrity_log(file_path, "Apply Anti-Debugging", "Applied")
    
            if args.detect_debug:
    
                insert_file_integrity_log(file_path, "Detect Anti-Debugging", "Detected" if detect_anti_debugging_and_obfuscation(file_path) else "Not Detected")
    
            if args.remove_debug:
    
                insert_file_integrity_log(file_path, "Remove Anti-Debugging", "Removed")

    # 네트워크 모니터링 옵션 처리
    if args.network_monitor:
    
        monitor_network()

    # DB 연결 종료
    if db_conn:
    
        db_conn.close()
    
        print("데이터베이스 연결이 종료되었습니다.")
        
    if args.report:
        
        process_report_option(args.report)

if __name__ == "__main__":
    
    main()

